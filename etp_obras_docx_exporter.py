from __future__ import annotations

from datetime import date
from decimal import Decimal
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.text.paragraph import Paragraph

from etp_obras import texto, validar_etp_obras

TEMPLATE = Path(__file__).resolve().parent / "templates" / "etp" / "obras_engenharia.docx"
MESES = (
    "janeiro", "fevereiro", "março", "abril", "maio", "junho",
    "julho", "agosto", "setembro", "outubro", "novembro", "dezembro",
)


def substituir(paragrafo: Paragraph, valor: str) -> None:
    if paragrafo.runs:
        paragrafo.runs[0].text = valor
        for trecho in paragrafo.runs[1:]:
            trecho.text = ""
    else:
        paragrafo.add_run(valor)


def reais(valor: Any) -> str:
    formatado = f"{Decimal(str(valor)):,.2f}"
    return "R$ " + formatado.replace(",", "X").replace(".", ",").replace("X", ".")


def gerar_etp_obras_docx(dados: dict[str, Any]) -> tuple[bytes, str]:
    validar_etp_obras(dados)
    if not TEMPLATE.is_file():
        raise ValueError("O modelo Word do ETP de obras não foi encontrado.")
    documento = Document(TEMPLATE)
    documento.core_properties.title = "ETP - Obras e Serviços de Engenharia"
    documento.core_properties.author = "Prefeitura Municipal de Mallet"
    documento.core_properties.last_modified_by = "Assistente de Contratações Públicas"

    valores = {
        4: f"OBJETO: {texto(dados['objeto'])}",
        5: f"SOLICITANTE: {texto(dados['solicitante'])}",
        6: f"RESPONSÁVEL PELA ELABORAÇÃO DO ETP: {texto(dados['responsavel_nome'])}",
        8: texto(dados["justificativa"]), 9: "",
        11: texto(dados["resultados_pretendidos"]),
        19: "Soluções identificadas no levantamento de mercado:",
        20: texto(dados["alternativas"]), 21: "",
        23: texto(dados["solucao_escolhida"]),
        39: f"Local de execução: {texto(dados['local_execucao'])}",
        40: f"Serviços e materiais: {texto(dados['servicos_materiais'])}",
        41: f"Metodologia executiva: {texto(dados['metodologia_executiva'])}",
        42: f"Cronograma físico-financeiro: {texto(dados['cronograma'])}",
        43: f"Objeto de engenharia: {texto(dados['objeto'])}",
        49: texto(dados["estimativa_quantidades"]), 50: "",
        52: (
            f"Fontes utilizadas: {texto(dados['fontes_pesquisa'])}\n"
            f"Metodologia: {texto(dados['metodologia_pesquisa'])}"
        ),
        53: "",
        55: texto(dados["justificativa_solucao"]),
        57: texto(dados["custos_operacionais"]),
        59: texto(dados["interferencias"]), 60: "",
        62: texto(dados["titularidade_area"]),
        65: f"Valor estimado da contratação: {reais(dados['valor_estimado'])}.",
        67: texto(dados["previsao_planos"]),
        69: texto(dados["parcelamento"]), 70: "", 71: "", 72: "", 73: "",
        76: texto(dados["contratacoes_correlatas"]),
        78: "", 79: "", 80: "",
        82: texto(dados["capacitacao"]), 83: "", 84: "",
        86: texto(dados["impactos_ambientais"]), 87: "",
        95: texto(dados["declaracao_viabilidade"]),
        119: texto(dados["anexos"]), 120: "",
    }
    for indice, valor in valores.items():
        substituir(documento.paragraphs[indice], valor)

    hoje = date.today()
    substituir(
        documento.paragraphs[122],
        f"Prefeitura Municipal de Mallet, {hoje.day} de {MESES[hoje.month - 1]} de {hoje.year}.",
    )
    assinatura = documento.tables[-1]
    assinatura.cell(0, 0).text = texto(dados["responsavel_nome"])
    assinatura.cell(1, 0).text = "RESPONSÁVEL PELA ELABORAÇÃO DO ETP"
    assinatura.cell(0, 1).text = texto(dados["autoridade_nome"])
    assinatura.cell(1, 1).text = texto(dados["autoridade_cargo"])

    arquivo = BytesIO()
    documento.save(arquivo)
    return arquivo.getvalue(), f"ETP_obras_engenharia_{hoje.strftime('%Y%m%d')}.docx"
