import unittest
from io import BytesIO
from zipfile import ZipFile

from docx import Document

from tr import TIPOS_TR, gerar_tr, listar_tipos_tr
from tr_docx_exporter import gerar_tr_docx


def dados_validos() -> dict:
    return {
        "secretaria": "Secretaria Municipal de Educação",
        "objeto": "Aquisição de materiais para a rede municipal.",
        "fundamentacao": "A contratação atende à necessidade demonstrada no ETP.",
        "metodologia_calculo": "Mediana de preços coletados no PNCP.",
        "vigencia_meses": 12,
        "criterio_selecao": "Menor preço por item.",
        "requisitos": "Produtos novos e em conformidade com as especificações.",
        "obrigacoes": "Entregar os itens nos prazos e locais definidos.",
        "recebimento": "Provisório em 5 dias e definitivo em 10 dias.",
        "pagamento": "Até 30 dias após o recebimento definitivo.",
        "reajuste": "INPC após o interregno legal.",
        "fiscal_nome": "Fiscal de teste",
        "fiscal_portaria": "Portaria nº 001/2027",
        "subcontratacao": "Não será permitida a subcontratação.",
        "previsao_orcamentaria": "Compatível com o PPA, a LDO e a LOA.",
        "dotacoes": "Dotação de teste 00.000.0000.0000.",
        "anexos": "ETP; pesquisa de preços; mapa de riscos.",
        "responsavel_nome": "Responsável de teste",
        "autoridade_nome": "Autoridade de teste",
        "autoridade_cargo": "Secretária Municipal de Educação",
        "fornecedor_nome": "Fornecedor de teste",
        "fornecedor_documento": "00.000.000/0001-00",
        "razoes_escolha": "Menor preço e atendimento aos requisitos.",
        "itens": [
            {
                "descricao": "Material de teste",
                "codigo": "1234",
                "quantidade": "10",
                "valor_unitario": "25.50",
            }
        ],
    }


class TestTR(unittest.TestCase):
    def test_lista_sete_modelos(self) -> None:
        self.assertEqual(len(listar_tipos_tr()), 7)

    def test_gera_previa(self) -> None:
        resultado = gerar_tr("compras", dados_validos())
        self.assertIn("TERMO DE REFERÊNCIA", resultado["conteudo"])
        self.assertIn("Menor preço", resultado["conteudo"])

    def test_exige_fornecedor_na_dispensa(self) -> None:
        dados = dados_validos()
        dados["fornecedor_nome"] = ""
        with self.assertRaisesRegex(ValueError, "Fornecedor escolhido"):
            gerar_tr("compras_dispensa_inexigibilidade", dados)

    def test_gera_os_sete_documentos_word(self) -> None:
        for tipo, config in TIPOS_TR.items():
            with self.subTest(tipo=tipo):
                dados = dados_validos()
                if not config["exige_itens"]:
                    dados["itens"] = []
                conteudo, nome = gerar_tr_docx(tipo, dados)
                self.assertTrue(nome.endswith(".docx"))
                with ZipFile(BytesIO(conteudo)) as arquivo:
                    self.assertIsNone(arquivo.testzip())
                documento = Document(BytesIO(conteudo))
                texto = "\n".join(p.text for p in documento.paragraphs)
                self.assertIn("Aquisição de materiais", texto)
                self.assertIn("Fiscal de teste", texto)
                self.assertEqual(
                    documento.tables[-1].cell(0, 1).text, "Autoridade de teste"
                )


if __name__ == "__main__":
    unittest.main()
