import unittest
from io import BytesIO
from zipfile import ZipFile

from docx import Document

from docx_exporter import gerar_dfd_docx
from test_dfd import dados_validos


class TestExportacaoDOCX(unittest.TestCase):
    def test_gera_docx_valido_a_partir_do_modelo_oficial(self) -> None:
        conteudo, nome = gerar_dfd_docx("pregao", dados_validos())
        self.assertTrue(nome.endswith(".docx"))
        self.assertGreater(len(conteudo), 10_000)
        with ZipFile(BytesIO(conteudo)) as arquivo:
            self.assertIsNone(arquivo.testzip())

    def test_insere_dados_no_documento(self) -> None:
        conteudo, _ = gerar_dfd_docx("pregao", dados_validos())
        documento = Document(BytesIO(conteudo))
        texto = "\n".join(paragrafo.text for paragrafo in documento.paragraphs)
        self.assertIn("Aquisição de materiais escolares", texto)
        self.assertIn("R$ 12.500,50", texto)
        self.assertIn("Autoridade de teste", texto)
        self.assertNotIn("OBJETO: xxxxxx", texto)

    def test_usa_modelo_especifico_para_dispensa(self) -> None:
        dados = dados_validos()
        dados["contratada_nome"] = "Empresa de Teste Ltda."
        dados["contratada_documento"] = "00.000.000/0001-00"
        conteudo, nome = gerar_dfd_docx("dispensa", dados)
        documento = Document(BytesIO(conteudo))
        texto = "\n".join(paragrafo.text for paragrafo in documento.paragraphs)
        self.assertIn("DISPENSA DE LICITAÇÃO", texto)
        self.assertIn("Empresa de Teste Ltda.", texto)
        self.assertIn("dispensa", nome)


if __name__ == "__main__":
    unittest.main()
