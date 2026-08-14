import unittest
from io import BytesIO
from zipfile import ZipFile

from docx import Document

from etp_obras import gerar_etp_obras
from etp_obras_docx_exporter import gerar_etp_obras_docx


def dados_validos() -> dict:
    return {
        "objeto": "Reforma da Escola Municipal de Teste",
        "solicitante": "Secretaria Municipal de Educação",
        "responsavel_nome": "Engenheiro de teste",
        "justificativa": "Corrigir problemas estruturais e garantir segurança.",
        "resultados_pretendidos": "Ambiente escolar seguro e acessível.",
        "alternativas": "Reforma integral ou intervenções parciais.",
        "solucao_escolhida": "Reforma integral por apresentar maior durabilidade.",
        "local_execucao": "Rua de Teste, 100, Mallet/PR.",
        "servicos_materiais": "Cobertura, instalações elétricas e pintura.",
        "metodologia_executiva": "Execução conforme projetos e normas técnicas.",
        "cronograma": "Execução estimada em seis meses.",
        "estimativa_quantidades": "Quantidades obtidas no levantamento técnico.",
        "fontes_pesquisa": "SINAPI e contratações similares.",
        "metodologia_pesquisa": "Composições do SINAPI vigente.",
        "justificativa_solucao": "Melhor relação entre custo e vida útil.",
        "custos_operacionais": "Manutenção preventiva anual.",
        "interferencias": "Rede elétrica existente identificada em vistoria.",
        "titularidade_area": "Imóvel registrado em nome do Município.",
        "valor_estimado": "250000.00",
        "previsao_planos": "Prevista no PPA e no PCA vigente.",
        "parcelamento": "Não recomendado devido à interdependência dos serviços.",
        "contratacoes_correlatas": "Aquisição posterior de mobiliário.",
        "capacitacao": "Fiscalização por profissional habilitado.",
        "impactos_ambientais": "Gestão e destinação adequada dos resíduos.",
        "declaracao_viabilidade": "A contratação é técnica e economicamente viável.",
        "anexos": "Projeto básico; memorial descritivo; levantamento fotográfico.",
        "autoridade_nome": "Autoridade de teste",
        "autoridade_cargo": "Secretária Municipal de Educação",
    }


class TestETPObras(unittest.TestCase):
    def test_gera_previa(self) -> None:
        resultado = gerar_etp_obras(dados_validos())
        self.assertIn("OBRAS E SERVIÇOS DE ENGENHARIA", resultado["conteudo"])

    def test_gera_word_oficial_preenchido(self) -> None:
        conteudo, nome = gerar_etp_obras_docx(dados_validos())
        self.assertTrue(nome.endswith(".docx"))
        with ZipFile(BytesIO(conteudo)) as arquivo:
            self.assertIsNone(arquivo.testzip())
        documento = Document(BytesIO(conteudo))
        texto = "\n".join(p.text for p in documento.paragraphs)
        self.assertIn("Reforma da Escola Municipal de Teste", texto)
        self.assertIn("R$ 250.000,00", texto)
        self.assertIn("SINAPI", texto)
        self.assertEqual(
            documento.tables[-1].cell(0, 1).text, "Autoridade de teste"
        )


if __name__ == "__main__":
    unittest.main()
