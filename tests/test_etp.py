import unittest
from io import BytesIO
from zipfile import ZipFile

from docx import Document

from etp import gerar_etp, validar_etp
from etp_docx_exporter import gerar_etp_docx


def dados_etp_validos() -> dict:
    return {
        "objeto": "Aquisição de kits de materiais escolares",
        "solicitante": "Secretaria Municipal de Educação",
        "responsavel_nome": "Responsável de teste",
        "justificativa": "Garantir materiais aos estudantes da rede municipal.",
        "resultados_pretendidos": "Atender os alunos com economicidade.",
        "alternativas": "Aquisição direta de kits ou aquisição de itens separados.",
        "solucao_escolhida": "Aquisição de kits completos.",
        "requisitos": "Materiais novos e em conformidade com as especificações.",
        "prazo_meses": 12,
        "itens": [
            {"descricao": "Kit de material escolar", "quantidade": "500 unidades"},
            {"descricao": "Agenda escolar", "quantidade": "500 unidades"},
        ],
        "memoria_calculo": "Quantidade baseada no número de matrículas.",
        "fontes_pesquisa": "PNCP e três orçamentos de fornecedores.",
        "metodologia_pesquisa": "Mediana dos valores válidos coletados.",
        "justificativa_solucao": "A solução apresenta melhor custo-benefício.",
        "valor_estimado": "35000.50",
        "previsao_planos": "Contratação prevista no PPA e no PCA vigente.",
        "parcelamento": "O objeto será parcelado por item para ampliar a competição.",
        "contratacoes_correlatas": "Não foram identificadas contratações dependentes.",
        "capacitacao": "Os fiscais designados possuem experiência compatível.",
        "impactos_ambientais": "Exigência de embalagens recicláveis.",
        "declaracao_viabilidade": "A contratação é técnica e economicamente viável.",
        "autoridade_nome": "Autoridade de teste",
        "autoridade_cargo": "Secretária Municipal de Educação",
    }


class TestETP(unittest.TestCase):
    def test_gera_previa(self) -> None:
        resultado = gerar_etp(dados_etp_validos())
        self.assertIn("COMPRAS E SERVIÇOS", resultado["conteudo"])
        self.assertIn("Garantir materiais", resultado["conteudo"])

    def test_exige_item(self) -> None:
        dados = dados_etp_validos()
        dados["itens"] = []
        with self.assertRaisesRegex(ValueError, "pelo menos um item"):
            validar_etp(dados)

    def test_gera_word_integro_e_preenchido(self) -> None:
        conteudo, nome = gerar_etp_docx(dados_etp_validos())
        self.assertTrue(nome.endswith(".docx"))
        with ZipFile(BytesIO(conteudo)) as arquivo:
            self.assertIsNone(arquivo.testzip())
        documento = Document(BytesIO(conteudo))
        texto = "\n".join(p.text for p in documento.paragraphs)
        self.assertIn("Aquisição de kits de materiais escolares", texto)
        self.assertIn("R$ 35.000,50", texto)
        tabela = documento.tables[0]
        self.assertEqual(len(tabela.rows), 3)
        self.assertEqual(tabela.cell(1, 1).text, "Kit de material escolar")
        self.assertEqual(
            documento.tables[-1].cell(0, 1).text, "Autoridade de teste"
        )


if __name__ == "__main__":
    unittest.main()
