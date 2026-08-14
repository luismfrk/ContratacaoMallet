from __future__ import annotations

import re
from datetime import date
from decimal import Decimal
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from tr import TIPOS_TR, texto, validar_tr

TEMPLATE_DIR = Path(__file__).resolve().parent / "templates" / "tr"
MESES = (
    "janeiro", "fevereiro", "março", "abril", "maio", "junho",
    "julho", "agosto", "setembro", "outubro", "novembro", "dezembro",
)


def _normalizar(valor: str) -> str:
    return re.sub(r"^\d+\.\s*", "", " ".join(valor.upper().split()))


def _inserir_depois(paragrafo: Paragraph, conteudo: str) -> Paragraph:
    novo_elemento = OxmlElement("w:p")
    paragrafo._p.addnext(novo_elemento)
    novo = Paragraph(novo_elemento, paragrafo._parent)
    novo.style = paragrafo.style
    novo.add_run(conteudo)
    return novo


def _reais(valor: Any) -> str:
    formatado = f"{Decimal(str(valor)):,.2f}"
    return "R$ " + formatado.replace(",", "X").replace(".", ",").replace("X", ".")


def _preencher_itens(documento: Document, itens: list[dict[str, Any]]) -> None:
    tabela = None
    for candidata in documento.tables:
        if not candidata.rows:
            continue
        cabecalho = " | ".join(c.text for c in candidata.rows[0].cells)
        if "Item" in cabecalho and "Descrição" in cabecalho and "Qtd." in cabecalho:
            tabela = candidata
            break
    if tabela is None:
        return
    while len(tabela.rows) > 1:
        tabela._tbl.remove(tabela.rows[-1]._tr)
    headers = [" ".join(c.text.split()).lower() for c in tabela.rows[0].cells]
    for indice, item in enumerate(itens, start=1):
        celulas = tabela.add_row().cells
        total = Decimal(str(item["quantidade"])) * Decimal(str(item["valor_unitario"]))
        for coluna, cabecalho in enumerate(headers):
            if cabecalho == "item":
                celulas[coluna].text = str(indice)
            elif "descrição" in cabecalho:
                celulas[coluna].text = texto(item["descricao"])
            elif "catser" in cabecalho:
                celulas[coluna].text = texto(item.get("codigo"))
            elif "qtd" in cabecalho:
                celulas[coluna].text = texto(item["quantidade"])
            elif "unit" in cabecalho or "mensal" in cabecalho:
                celulas[coluna].text = _reais(item["valor_unitario"])
            elif "total" in cabecalho:
                celulas[coluna].text = _reais(total)


def _preencher_tabelas_administrativas(documento: Document, dados: dict[str, Any]) -> None:
    for tabela in documento.tables:
        if not tabela.rows:
            continue
        cabecalho = " | ".join(" ".join(c.text.split()) for c in tabela.rows[0].cells)
        if "Informações gerais da contratação" in cabecalho:
            for linha in tabela.rows:
                rotulo = linha.cells[0].text.upper()
                if "SECRETARIA" in rotulo:
                    linha.cells[-1].text = texto(dados["secretaria"])
                elif "EMPRESA" in rotulo:
                    linha.cells[-1].text = texto(dados.get("fornecedor_nome"))
                elif "CNPJ" in rotulo or "CPF" in rotulo:
                    linha.cells[-1].text = texto(dados.get("fornecedor_documento"))
                elif "VALOR" in rotulo:
                    total = sum(
                        Decimal(str(i["quantidade"])) * Decimal(str(i["valor_unitario"]))
                        for i in dados.get("itens", [])
                    )
                    linha.cells[-1].text = _reais(total)
        elif "Secretaria | Fiscal | Portaria" in cabecalho and len(tabela.rows) > 1:
            tabela.cell(1, 0).text = texto(dados["secretaria"])
            tabela.cell(1, 1).text = texto(dados["fiscal_nome"])
            tabela.cell(1, 2).text = texto(dados["fiscal_portaria"])

    assinatura = documento.tables[-1]
    if len(assinatura.rows) >= 2 and len(assinatura.columns) >= 2:
        assinatura.cell(0, 0).text = texto(dados["responsavel_nome"])
        assinatura.cell(1, 0).text = "RESPONSÁVEL PELA ELABORAÇÃO DO TR"
        assinatura.cell(0, 1).text = texto(dados["autoridade_nome"])
        assinatura.cell(1, 1).text = texto(dados["autoridade_cargo"])


def gerar_tr_docx(tipo: str, dados: dict[str, Any]) -> tuple[bytes, str]:
    validar_tr(tipo, dados)
    modelo = TEMPLATE_DIR / f"{tipo}.docx"
    if not modelo.is_file():
        raise ValueError("O modelo Word desse Termo de Referência não foi encontrado.")
    documento = Document(modelo)
    documento.core_properties.title = f"TR - {TIPOS_TR[tipo]['titulo']}"
    documento.core_properties.author = "Prefeitura Municipal de Mallet"
    documento.core_properties.last_modified_by = "Assistente de Contratações Públicas"

    conteudos = (
        ("FUNDAMENTAÇÃO", texto(dados["fundamentacao"])),
        ("DO OBJETO", texto(dados["objeto"])),
        ("METODOLOGIA DE CÁLCULO", texto(dados["metodologia_calculo"])),
        ("PRAZO DE VIGÊNCIA", f"{int(dados['vigencia_meses'])} meses."),
        ("PRAZOS DE VIGÊNCIA", f"Vigência: {int(dados['vigencia_meses'])} meses."),
        ("FORMA E CRITÉRIOS DE SELEÇÃO", texto(dados["criterio_selecao"])),
        ("REQUISITOS DA CONTRATAÇÃO", texto(dados["requisitos"])),
        ("OBRIGAÇÕES DA CONTRATADA", texto(dados["obrigacoes"])),
        ("RECEBIMENTO PROVISÓRIO", texto(dados["recebimento"])),
        ("CONDIÇÕES PARA PAGAMENTO", texto(dados["pagamento"])),
        ("DO REAJUSTE", texto(dados["reajuste"])),
        (
            "FISCALIZAÇÃO DO CONTRATO",
            f"Fiscal: {texto(dados['fiscal_nome'])}. Portaria: {texto(dados['fiscal_portaria'])}.",
        ),
        ("SUBCONTRATAÇÃO", texto(dados["subcontratacao"])),
        (
            "PREVISÃO ORÇAMENTÁRIA",
            f"{texto(dados['previsao_orcamentaria'])}\nDotações: {texto(dados['dotacoes'])}",
        ),
        ("DOS ANEXOS", texto(dados["anexos"])),
        ("EMPRESA VENCEDORA", f"{texto(dados.get('fornecedor_nome'))} – {texto(dados.get('fornecedor_documento'))}"),
        ("EMPRESA ESCOLHIDA", f"{texto(dados.get('fornecedor_nome'))} – {texto(dados.get('fornecedor_documento'))}"),
        ("RAZÕES PARA ESCOLHA", texto(dados.get("razoes_escolha"))),
    )
    usados: set[str] = set()
    for paragrafo in list(documento.paragraphs):
        titulo = _normalizar(paragrafo.text)
        for chave, valor in conteudos:
            if chave not in usados and chave in titulo and valor:
                _inserir_depois(paragrafo, valor)
                usados.add(chave)
                break

    _preencher_itens(documento, dados.get("itens", []))
    _preencher_tabelas_administrativas(documento, dados)
    hoje = date.today()
    for paragrafo in documento.paragraphs:
        if paragrafo.text.strip().startswith("Prefeitura Municipal de Mallet,"):
            paragrafo.text = (
                f"Prefeitura Municipal de Mallet, {hoje.day} de "
                f"{MESES[hoje.month - 1]} de {hoje.year}."
            )

    arquivo = BytesIO()
    documento.save(arquivo)
    return arquivo.getvalue(), f"TR_{tipo}_{hoje.strftime('%Y%m%d')}.docx"
