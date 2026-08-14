from __future__ import annotations

from datetime import date
from decimal import Decimal
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.text.paragraph import Paragraph

from app import TIPOS_DFD, validar_dfd

BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_DIR = BASE_DIR / "templates" / "dfd"
MESES = (
    "janeiro",
    "fevereiro",
    "março",
    "abril",
    "maio",
    "junho",
    "julho",
    "agosto",
    "setembro",
    "outubro",
    "novembro",
    "dezembro",
)


def _texto(valor: Any) -> str:
    if valor is None:
        return ""
    return str(valor).strip()


def _valor_em_reais(valor: Any) -> str:
    numero = Decimal(str(valor))
    formatado = f"{numero:,.2f}"
    return "R$ " + formatado.replace(",", "X").replace(".", ",").replace("X", ".")


def _substituir_texto(paragrafo: Paragraph, texto: str) -> None:
    if paragrafo.runs:
        paragrafo.runs[0].text = texto
        for trecho in paragrafo.runs[1:]:
            trecho.text = ""
    else:
        paragrafo.add_run(texto)


def _data_por_extenso(valor: date) -> str:
    return f"Mallet, {valor.day} de {MESES[valor.month - 1]} de {valor.year}."


def gerar_dfd_docx(tipo: str, dados: dict[str, Any]) -> tuple[bytes, str]:
    validar_dfd(tipo, dados)
    configuracao = TIPOS_DFD[tipo]
    caminho_modelo = TEMPLATE_DIR / f"{tipo}.docx"
    if not caminho_modelo.is_file():
        raise ValueError(f"O modelo Word para {configuracao.titulo} não foi encontrado.")

    documento = Document(caminho_modelo)
    documento.core_properties.title = f"DFD - {configuracao.titulo}"
    documento.core_properties.subject = "Documento de Formalização de Demanda"
    documento.core_properties.author = "Prefeitura Municipal de Mallet"
    documento.core_properties.last_modified_by = "Assistente de Contratações Públicas"

    proximo_conteudo: str | None = None
    for paragrafo in documento.paragraphs:
        atual = " ".join(paragrafo.text.split())

        if proximo_conteudo is not None and atual.startswith("Todos os DFDs"):
            _substituir_texto(paragrafo, proximo_conteudo)
            proximo_conteudo = None
        elif atual.startswith("IDENTIFICAÇÃO DA UNIDADE REQUISITANTE:"):
            _substituir_texto(
                paragrafo,
                f"IDENTIFICAÇÃO DA UNIDADE REQUISITANTE: {_texto(dados['unidade_requisitante'])}",
            )
        elif atual.startswith("OBJETO:"):
            _substituir_texto(paragrafo, f"OBJETO: {_texto(dados['objeto'])}")
        elif atual == "JUSTIFICATIVA DA NECESSIDADE DA CONTRATAÇÃO:":
            proximo_conteudo = _texto(dados["justificativa"])
        elif atual.startswith("QUANTITATIVO:"):
            _substituir_texto(
                paragrafo,
                f"QUANTITATIVO: {_texto(dados['quantitativo'])} {_texto(dados['unidade_medida'])}",
            )
            proximo_conteudo = ""
        elif atual.startswith("VALOR INICIAL:"):
            _substituir_texto(
                paragrafo, f"VALOR INICIAL: {_valor_em_reais(dados['valor_estimado'])}"
            )
        elif atual.startswith("Todos os DFDs do ano de referência devem apresentar"):
            _substituir_texto(
                paragrafo, f"FONTE DA ESTIMATIVA: {_texto(dados['fonte_estimativa'])}"
            )
        elif atual.startswith("PRAZO DE VIGÊNCIA DO CONTRATO:"):
            _substituir_texto(
                paragrafo,
                f"PRAZO DE VIGÊNCIA DO CONTRATO: {int(dados['prazo_vigencia_meses'])} meses.",
            )
        elif atual.startswith("DATA PRETENDIDA DA CONTRATAÇÃO:"):
            data_pretendida = date.fromisoformat(_texto(dados["data_pretendida"]))
            _substituir_texto(
                paragrafo,
                f"DATA PRETENDIDA DA CONTRATAÇÃO: {data_pretendida.strftime('%d/%m/%Y')}",
            )
        elif atual.startswith("CONTRATADA:") and configuracao.exige_contratada:
            _substituir_texto(
                paragrafo,
                f"CONTRATADA: {_texto(dados['contratada_nome'])} "
                f"({_texto(dados['contratada_documento'])})",
            )
        elif atual.startswith("RESPONSÁVEL PELA SOLICITAÇÃO:"):
            _substituir_texto(
                paragrafo,
                "RESPONSÁVEL PELA SOLICITAÇÃO: "
                f"{_texto(dados['responsavel_nome'])}, "
                f"{_texto(dados['responsavel_cargo'])}, matrícula "
                f"{_texto(dados['responsavel_matricula'])}",
            )
        elif atual.startswith("CÓDIGO PNCP:"):
            _substituir_texto(
                paragrafo,
                f"CÓDIGO PNCP: {_texto(dados.get('codigo_pncp')) or 'Não informado'}",
            )
        elif atual.startswith("FUNDAMENTO:"):
            _substituir_texto(paragrafo, f"FUNDAMENTO: {_texto(dados['fundamento'])}")
        elif atual.startswith("MODALIDADE:"):
            _substituir_texto(paragrafo, f"MODALIDADE: {configuracao.titulo}")
        elif atual.startswith("Mallet,"):
            _substituir_texto(paragrafo, _data_por_extenso(date.today()))
        elif atual == "NOME":
            _substituir_texto(paragrafo, _texto(dados["autoridade_nome"]))
        elif atual == "Secretário Municipal de xxxxx":
            _substituir_texto(paragrafo, _texto(dados["autoridade_cargo"]))

    arquivo = BytesIO()
    documento.save(arquivo)
    nome_arquivo = f"DFD_{tipo}_{date.today().strftime('%Y%m%d')}.docx"
    return arquivo.getvalue(), nome_arquivo
