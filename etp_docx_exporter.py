from __future__ import annotations

from datetime import date
from decimal import Decimal
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.text.paragraph import Paragraph

from etp import validar_etp

BASE_DIR = Path(__file__).resolve().parent
TEMPLATE = BASE_DIR / "templates" / "etp" / "compras_servicos.docx"
MESES = (
    "janeiro", "fevereiro", "março", "abril", "maio", "junho",
    "julho", "agosto", "setembro", "outubro", "novembro", "dezembro",
)


def _texto(valor: Any) -> str:
    return "" if valor is None else str(valor).strip()


def _substituir(paragrafo: Paragraph, texto: str) -> None:
    if paragrafo.runs:
        paragrafo.runs[0].text = texto
        for trecho in paragrafo.runs[1:]:
            trecho.text = ""
    else:
        paragrafo.add_run(texto)


def _reais(valor: Any) -> str:
    formatado = f"{Decimal(str(valor)):,.2f}"
    return "R$ " + formatado.replace(",", "X").replace(".", ",").replace("X", ".")


def _preencher_tabela_itens(documento: Document, itens: list[dict[str, Any]]) -> None:
    tabela = documento.tables[0]
    while len(tabela.rows) > 1:
        tabela._tbl.remove(tabela.rows[-1]._tr)
    for indice, item in enumerate(itens, start=1):
        celulas = tabela.add_row().cells
        celulas[0].text = str(indice)
        celulas[1].text = _texto(item["descricao"])
        celulas[2].text = _texto(item["quantidade"])


def gerar_etp_docx(dados: dict[str, Any]) -> tuple[bytes, str]:
    validar_etp(dados)
    if not TEMPLATE.is_file():
        raise ValueError("O modelo Word do ETP de compras e serviços não foi encontrado.")

    documento = Document(TEMPLATE)
    documento.core_properties.title = "ETP - Compras e Serviços"
    documento.core_properties.author = "Prefeitura Municipal de Mallet"
    documento.core_properties.last_modified_by = "Assistente de Contratações Públicas"

    substituicoes = {
        4: f"OBJETO: {_texto(dados['objeto'])}",
        5: f"SOLICITANTE: {_texto(dados['solicitante'])}",
        6: f"RESPONSÁVEL PELA ELABORAÇÃO DO ETP: {_texto(dados['responsavel_nome'])}",
        8: _texto(dados["justificativa"]),
        9: "",
        11: _texto(dados["resultados_pretendidos"]),
        18: "Soluções identificadas no levantamento de mercado:",
        19: _texto(dados["alternativas"]),
        20: "",
        22: _texto(dados["solucao_escolhida"]),
        29: _texto(dados["requisitos"]),
        30: (
            f"O prazo do contrato referente ao objeto deste Estudo Técnico será de "
            f"{int(dados['prazo_meses'])} meses. Os demais requisitos serão definidos "
            "no Termo de Referência."
        ),
        33: _texto(dados["memoria_calculo"]),
        34: "",
        36: (
            f"Fontes utilizadas: {_texto(dados['fontes_pesquisa'])}\n"
            f"Metodologia adotada: {_texto(dados['metodologia_pesquisa'])}"
        ),
        37: "",
        41: _texto(dados["justificativa_solucao"]),
        43: f"O custo total estimado da contratação é de {_reais(dados['valor_estimado'])}.",
        46: _texto(dados["previsao_planos"]),
        48: _texto(dados["parcelamento"]),
        50: _texto(dados["contratacoes_correlatas"]),
        53: _texto(dados["capacitacao"]),
        55: _texto(dados["impactos_ambientais"]),
        64: _texto(dados["declaracao_viabilidade"]),
    }
    for indice, texto in substituicoes.items():
        _substituir(documento.paragraphs[indice], texto)

    hoje = date.today()
    _substituir(
        documento.paragraphs[66],
        f"Prefeitura Municipal de Mallet, {hoje.day} de {MESES[hoje.month - 1]} de {hoje.year}.",
    )
    _preencher_tabela_itens(documento, dados["itens"])

    assinatura = documento.tables[-1]
    assinatura.cell(0, 0).text = _texto(dados["responsavel_nome"])
    assinatura.cell(1, 0).text = "RESPONSÁVEL PELA ELABORAÇÃO DO ETP"
    assinatura.cell(0, 1).text = _texto(dados["autoridade_nome"])
    assinatura.cell(1, 1).text = _texto(dados["autoridade_cargo"])

    arquivo = BytesIO()
    documento.save(arquivo)
    return arquivo.getvalue(), f"ETP_compras_servicos_{hoje.strftime('%Y%m%d')}.docx"
